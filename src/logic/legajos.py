import os
import logging
from typing import Dict, List, Any
# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx2pdf import convert

# Configuración básica de logging para monitoreo de batch processing
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

def _replace_in_paragraphs(paragraphs: List[Any], data: Dict[str, Any], doc: Document = None) -> None:
    """
    Función auxiliar para reemplazar texto a nivel de párrafo.
    Resuelve el problema de los marcadores divididos en múltiples 'runs'
    y preserva el formato original del texto (negritas, cursivas, etc.).
    """
    for paragraph in paragraphs:
        p_text = paragraph.text
        # Verificación rápida
        if not any(f"{{{{ {k} }}}}" in p_text or f"{{{{{k}}}}}" in p_text for k in data):
            continue

        # Lista de listas de caracteres para cada run
        run_chars = [list(run.text) for run in paragraph.runs]
        
        for key, value in data.items():
            is_table = isinstance(value, dict) and value.get("__type__") == "table"
            
            for placeholder in (f"{{{{ {key} }}}}", f"{{{{{key}}}}}"):
                while True:
                    current_text = "".join("".join(chars) for chars in run_chars)
                    start_idx = current_text.find(placeholder)
                    if start_idx == -1:
                        break
                        
                    if is_table and doc is not None:
                        # Crear e insertar tabla
                        table = doc.add_table(rows=1, cols=len(value["headers"]))
                        table.style = 'Table Grid'
                        table.autofit = True
                        table.allow_autofit = True
                        
                        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        from docx.shared import Pt, RGBColor, Cm
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        
                        # Configurar el encabezado para que se repita en nuevas páginas
                        tr = table.rows[0]._tr
                        trPr = tr.get_or_add_trPr()
                        tblHeader = OxmlElement('w:tblHeader')
                        tblHeader.set(qn('w:val'), "true")
                        trPr.append(tblHeader)
                        
                        # Set exact height for header row
                        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                        table.rows[0].height = Pt(12)

                        hdr_cells = table.rows[0].cells
                        for i, hdr in enumerate(value["headers"]):
                            hdr_cells[i].text = hdr
                            tcPr = hdr_cells[i]._element.get_or_add_tcPr()
                            
                            # Prevent wrapping
                            noWrap = OxmlElement('w:noWrap')
                            tcPr.append(noWrap)
                            
                            # Header Background
                            shd = tcPr.find(qn('w:shd'))
                            if shd is None:
                                shd = OxmlElement('w:shd')
                                tcPr.append(shd)
                            shd.set(qn('w:fill'), "4F81BD")  # Blue color
                            
                            # Cell margins (minimize to 20 dxa to fit text)
                            tcMar = OxmlElement('w:tcMar')
                            for margin in ['left', 'right']:
                                m = OxmlElement(f'w:{margin}')
                                m.set(qn('w:w'), "20")
                                m.set(qn('w:type'), "dxa")
                                tcMar.append(m)
                            tcPr.append(tcMar)
                            
                            # Font styling for header
                            for cell_p in hdr_cells[i].paragraphs:
                                cell_p.paragraph_format.space_after = Pt(2)
                                for run in cell_p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.size = Pt(5)
                                    
                        for r_idx, r_data in enumerate(value["rows"]):
                            new_row = table.add_row()
                            new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                            new_row.height = Pt(12)
                            row_cells = new_row.cells
                            is_even = (r_idx % 2 == 1)
                            is_totals = (r_idx == len(value["rows"]) - 1)
                            
                            for i, hdr in enumerate(value["headers"]):
                                row_cells[i].text = str(r_data.get(hdr, ""))
                                
                                tcPr = row_cells[i]._element.get_or_add_tcPr()
                                
                                # Prevent wrapping
                                noWrap = OxmlElement('w:noWrap')
                                tcPr.append(noWrap)
                                
                                # Alternating background
                                if is_even and not is_totals:
                                    shd = tcPr.find(qn('w:shd'))
                                    if shd is None:
                                        shd = OxmlElement('w:shd')
                                        tcPr.append(shd)
                                    shd.set(qn('w:fill'), "DCE6F1")  # Light blue
                                
                                # Cell margins (minimize)
                                tcMar = OxmlElement('w:tcMar')
                                for margin in ['left', 'right']:
                                    m = OxmlElement(f'w:{margin}')
                                    m.set(qn('w:w'), "20")
                                    m.set(qn('w:type'), "dxa")
                                    tcMar.append(m)
                                tcPr.append(tcMar)
                                
                                # Font size for data
                                for cell_p in row_cells[i].paragraphs:
                                    cell_p.paragraph_format.space_after = Pt(2)
                                    for run in cell_p.runs:
                                        run.font.size = Pt(5)
                                        if is_totals:
                                            run.font.bold = True
                        
                        # Mover la tabla justo después del párrafo actual
                        paragraph._p.addnext(table._tbl)
                        replace_val = ""
                    else:
                        replace_val = str(value)
                        
                    # Encontrar en qué run empieza el marcador
                    curr_len = 0
                    run_start = 0
                    char_start = 0
                    for r_idx, chars in enumerate(run_chars):
                        if curr_len <= start_idx < curr_len + len(chars):
                            run_start = r_idx
                            char_start = start_idx - curr_len
                            break
                        curr_len += len(chars)
                        
                    chars_to_remove = len(placeholder)
                    removable_in_start = len(run_chars[run_start]) - char_start
                    
                    if chars_to_remove <= removable_in_start:
                        # El marcador completo está en este run
                        run_chars[run_start][char_start:char_start+chars_to_remove] = list(replace_val)
                    else:
                        # El marcador abarca varios runs
                        run_chars[run_start] = run_chars[run_start][:char_start] + list(replace_val)
                        chars_to_remove -= removable_in_start
                        
                        r_idx = run_start + 1
                        while chars_to_remove > 0 and r_idx < len(run_chars):
                            removable_in_current = len(run_chars[r_idx])
                            if chars_to_remove >= removable_in_current:
                                run_chars[r_idx] = []
                                chars_to_remove -= removable_in_current
                            else:
                                run_chars[r_idx] = run_chars[r_idx][chars_to_remove:]
                                chars_to_remove = 0
                            r_idx += 1

        # Volcar los caracteres de nuevo a los runs preservando saltos de página (w:br)
        for run, chars in zip(paragraph.runs, run_chars):
            new_text = "".join(chars)
            if run.text != new_text:
                t_elements = run._element.xpath('.//w:t')
                if t_elements:
                    t_elements[0].text = new_text
                    for t in t_elements[1:]:
                        run._element.remove(t)
                else:
                    run.text = new_text

def replace_placeholders_in_doc(doc: Document, data: Dict[str, Any]) -> None:
    """
    Reemplaza los marcadores de posición {{ variable }} en un documento Word.

    Itera sobre los párrafos principales, el contenido de todas las tablas,
    y también sobre los encabezados y pies de página.

    Args:
        doc (Document): Objeto Document de python-docx ya instanciado.
        data (Dict[str, Any]): Diccionario con las claves a buscar y sus valores de reemplazo.

    Raises:
        TypeError: Si los tipos de datos en el diccionario no son convertibles a string.
    """
    def _process_blocks(blocks):
        if hasattr(blocks, 'paragraphs'):
            _replace_in_paragraphs(blocks.paragraphs, data, doc)
        if hasattr(blocks, 'tables'):
            for table in blocks.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_paragraphs(cell.paragraphs, data, doc)

    # 1. Reemplazo en cuerpo principal y tablas del cuerpo
    _process_blocks(doc)

    # 2. Reemplazo en encabezados y pies de página
    for section in doc.sections:
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, header_type, None)
            if header:
                _process_blocks(header)
                
        for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_type, None)
            if footer:
                _process_blocks(footer)

def process_docx(template_path: str, output_docx: str, data: Dict[str, Any]) -> None:
    """
    Carga una plantilla, reemplaza las variables estipuladas y
    lo exporta a formato .docx.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"No se encontró la plantilla en: {template_path}")
        
    try:
        from docx import Document
        doc = Document(template_path)
        replace_placeholders_in_doc(doc, data)
        doc.save(output_docx)
    except Exception as e:
        raise RuntimeError(f"Error interno modificando el documento .docx: {e}")

def process_document(template_path: str, output_pdf: str, data: Dict[str, Any]) -> None:
    """
    Carga una plantilla, reemplaza las variables estipuladas y
    lo exporta automáticamente a formato .pdf sin guardar el .docx modificado en disco de forma permanente.
    
    Args:
        template_path (str): Ruta absoluta o relativa al archivo .docx de origen.
        output_pdf (str): Ruta donde se exportará el archivo PDF final.
        data (Dict[str, Any]): Estructura de datos para reemplazar en el documento.

    Raises:
        FileNotFoundError: Si el archivo de la plantilla no existe en la ruta dada.
        RuntimeError: Ante cualquier falla inesperada del motor de conversión docx2pdf.
    """
    import tempfile
    
    temp_docx_fd, temp_docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(temp_docx_fd)
    
    # Fase 1: Procesamiento del .docx
    process_docx(template_path, temp_docx_path, data)

    # Fase 2: Conversión a PDF mediante COM/Word nativo (Windows) o LibreOffice (Linux)
    try:
        abs_docx = os.path.abspath(temp_docx_path)
        abs_pdf = os.path.abspath(output_pdf)
        
        import sys
        if sys.platform == "win32":
            convert(abs_docx, abs_pdf)
        else:
            import subprocess
            import shutil
            out_dir = os.path.dirname(abs_pdf)
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf", 
                "--outdir", out_dir, abs_docx
            ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            base_name = os.path.splitext(os.path.basename(abs_docx))[0]
            generated_pdf = os.path.join(out_dir, f"{base_name}.pdf")
            if os.path.exists(generated_pdf):
                shutil.move(generated_pdf, abs_pdf)
            else:
                raise RuntimeError("LibreOffice no generó el archivo PDF esperado.")
    except Exception as e:
        raise RuntimeError(f"Error durante la conversión a PDF: {e}")
    finally:
        # Limpiar el archivo temporal
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except OSError:
                pass

def batch_process_documents(template_path: str, output_dir: str, data_list: List[Dict[str, Any]]) -> None:
    """
    Procesa un lote completo de documentos a partir de un iterable de diccionarios, 
    gestionando las excepciones a nivel de iteración para no interrumpir el pipeline.

    Args:
        template_path (str): Ruta a la plantilla de Word original.
        output_dir (str): Directorio de salida para los archivos generados.
        data_list (List[Dict[str, Any]]): Lista de diccionarios (cada registro del DataFrame o JSON).
    
    Raises:
        ValueError: Si la lista de datos proporcionada está vacía.
    """
    if not data_list:
        raise ValueError("El lote de datos a procesar está vacío.")
    
    os.makedirs(output_dir, exist_ok=True)

    for idx, data in enumerate(data_list):
        # Utiliza una clave única si existe, de lo contrario un índice secuencial
        doc_name = str(data.get('id', f"documento_{idx}"))
        out_pdf = os.path.join(output_dir, f"{doc_name}.pdf")
        
        try:
            process_document(template_path, out_pdf, data)
            logging.info(f"Éxito: {doc_name}.pdf generado.")
        except Exception as e:
            logging.error(f"Fallo en {doc_name}: {e}")