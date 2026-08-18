import os
import shutil
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
import numpy_financial as npf

def monto_a_letras(monto: float) -> str:
    try:
        from num2words import num2words
        if monto is None:
            return ""
        monto_redondeado = round(float(monto), 2)
        entero = int(monto_redondeado)
        decimal = int(round((monto_redondeado - entero) * 100))
        letras_entero = num2words(entero, lang='es').upper()
        return f"{letras_entero} PESOS CON {decimal:02d}/100"
    except Exception as e:
        import traceback
        error_msg = f"Error en monto_a_letras con valor {monto}: {e}\n{traceback.format_exc()}"
        with open("error_monto.txt", "a") as f:
            f.write(error_msg + "\n")
        print(error_msg)
        return ""

def format_currency(monto: float) -> str:
    if monto is None:
        return "$ 0,00"
    try:
        formatted = f"{float(monto):,.2f}"
        # Swap commas and periods for Spanish locale
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"$ {formatted}"
    except Exception:
        return "$ 0,00"

def calcular_gastos_credito(credito) -> float:
    gastos = 0.0
    if getattr(credito, 'comision', None):
        gasto_1 = float(credito.comision.gasto_1_porcentaje or 0.0)
        gasto_2 = float(credito.comision.gasto_2_porcentaje or 0.0)
        sellado = float(getattr(credito.comision, 'porcentaje_sellado', 0) or 0.0)
        gastos = float(credito.capital) * (gasto_1 + gasto_2 + sellado) / 100.0
    else:
        # Fallback para creditos originados antes del parche de comision_id
        if getattr(credito, 'transferencias', None):
            gastos = sum(float(t.monto) for t in credito.transferencias if t.razon_social in ["AMUF", "DALVI CULTURAL SRL"])
    return gastos

from src.database import get_db, SocioComercial, Credito, Cliente, Cartera
from src.database.models.papeleria import DocumentoPapeleria, DocumentoVariable
from src.config import get_company_data

class VariableItem(BaseModel):
    placeholder: str
    system_field: str

class VariablesRequest(BaseModel):
    variables: List[VariableItem]

class ReorderItem(BaseModel):
    id: int
    orden: int

class ReorderRequest(BaseModel):
    documentos: List[ReorderItem]

SYSTEM_FIELDS = [
    {"value": "cliente.nombre", "label": "Cliente - Nombre Completo"},
    {"value": "cliente.apellido", "label": "Cliente - Apellido"},
    {"value": "cliente.dni", "label": "Cliente - DNI"},
    {"value": "cliente.cuil", "label": "Cliente - CUIL"},
    {"value": "cliente.id", "label": "Cliente - ID"},
    {"value": "cliente.domicilio", "label": "Cliente - Domicilio"},
    {"value": "cliente.calle", "label": "Cliente - Calle"},
    {"value": "cliente.calle_nro", "label": "Cliente - Número de Calle"},
    {"value": "cliente.piso", "label": "Cliente - Piso"},
    {"value": "cliente.depto", "label": "Cliente - Departamento"},
    {"value": "cliente.cp", "label": "Cliente - Código Postal"},
    {"value": "cliente.cbu", "label": "Cliente - CBU"},
    {"value": "cliente.cuenta_bancaria", "label": "Cliente - Cuenta Bancaria"},
    {"value": "cliente.banco", "label": "Cliente - Banco"},
    {"value": "cliente.localidad", "label": "Cliente - Localidad"},
    {"value": "cliente.provincia", "label": "Cliente - Provincia"},
    {"value": "cliente.nacionalidad", "label": "Cliente - País / Nacionalidad"},
    {"value": "cliente.estado_civil", "label": "Cliente - Estado Civil"},
    {"value": "cliente.telefono", "label": "Cliente - Teléfono"},
    {"value": "cliente.telefono_2", "label": "Cliente - Teléfono 2"},
    {"value": "cliente.mail", "label": "Cliente - Email"},
    {"value": "cliente.sexo", "label": "Cliente - Sexo"},
    {"value": "cliente.fecha_nacimiento", "label": "Cliente - Fecha Nacimiento"},
    {"value": "cliente.cargo", "label": "Cliente - Cargo"},
    {"value": "cliente.pep", "label": "Cliente - Es PEP"},
    {"value": "cliente.repet", "label": "Cliente - En REPET"},
    {"value": "cliente.repet_id", "label": "Cliente - ID RePET"},
    {"value": "cliente.repet_fecha_consulta", "label": "Cliente - Fecha y Hora Consulta RePET"},
    {"value": "cliente.repet_estado", "label": "Cliente - Estado RePET"},
    {"value": "referido_1.nombre", "label": "Referido 1 - Nombre"},
    {"value": "referido_1.apellido", "label": "Referido 1 - Apellido"},
    {"value": "referido_1.telefono", "label": "Referido 1 - Teléfono"},
    {"value": "referido_1.email", "label": "Referido 1 - Email"},
    {"value": "referido_2.nombre", "label": "Referido 2 - Nombre"},
    {"value": "referido_2.apellido", "label": "Referido 2 - Apellido"},
    {"value": "referido_2.telefono", "label": "Referido 2 - Teléfono"},
    {"value": "referido_2.email", "label": "Referido 2 - Email"},
    {"value": "empleador.razon_social", "label": "Empleador - Nombre / Razón Social"},
    {"value": "empleador.fecha_ingreso", "label": "Empleador - Fecha de Ingreso"},
    {"value": "empleador.ingreso_mensual", "label": "Empleador - Ingreso Mensual"},
    {"value": "empleador.legajo", "label": "Empleador - Legajo"},
    {"value": "empleador.domicilio_calle", "label": "Empleador - Domicilio Calle"},
    {"value": "empleador.domicilio_nro", "label": "Empleador - Domicilio Nro"},
    {"value": "empleador.domicilio_piso", "label": "Empleador - Domicilio Piso"},
    {"value": "empleador.domicilio_depto", "label": "Empleador - Domicilio Depto"},
    {"value": "empleador.localidad", "label": "Empleador - Localidad"},
    {"value": "empleador.provincia", "label": "Empleador - Provincia"},
    {"value": "empleador.codigo_postal", "label": "Empleador - Código Postal"},
    {"value": "empleador.telefono", "label": "Empleador - Teléfono"},
    {"value": "credito.monto_otorgado", "label": "Crédito - Monto Otorgado"},
    {"value": "credito.monto_otorgado_letras", "label": "Crédito - Monto Otorgado (En Letras)"},
    {"value": "credito.monto_neto", "label": "Crédito - Monto Neto (Otorgado - Transferencias)"},
    {"value": "credito.monto_neto_letras", "label": "Crédito - Monto Neto (En Letras)"},
    {"value": "credito.gastos_otorgamiento", "label": "Crédito - Gastos de Otorgamiento"},
    {"value": "credito.gastos_otorgamiento_letras", "label": "Crédito - Gastos de Otorgamiento (En Letras)"},
    {"value": "credito.sellado", "label": "Crédito - Sellado"},
    {"value": "credito.sellado_letras", "label": "Crédito - Sellado (En Letras)"},
    {"value": "credito.interes", "label": "Crédito - Interés Total"},
    {"value": "credito.interes_letras", "label": "Crédito - Interés Total (En Letras)"},
    {"value": "credito.iva", "label": "Crédito - IVA Total"},
    {"value": "credito.iva_letras", "label": "Crédito - IVA Total (En Letras)"},
    {"value": "credito.id", "label": "Crédito - ID Interno"},
    {"value": "credito.id_externo", "label": "Crédito - ID Externo"},
    {"value": "credito.plazo", "label": "Crédito - Plazo"},
    {"value": "credito.valor_cuota", "label": "Crédito - Valor Cuota"},
    {"value": "credito.monto_total", "label": "Crédito - Monto Total a Pagar"},
    {"value": "credito.monto_total_letras", "label": "Crédito - Monto Total a Pagar (En Letras)"},
    {"value": "credito.tna_c_iva", "label": "Crédito - TNA con IVA"},
    {"value": "credito.fecha_alta", "label": "Crédito - Fecha Alta"},
    {"value": "credito.fecha_emision_dia", "label": "Crédito - Fecha Emisión (Día)"},
    {"value": "credito.fecha_emision_mes_letras", "label": "Crédito - Fecha Emisión (Mes Letras)"},
    {"value": "credito.fecha_emision_anio", "label": "Crédito - Fecha Emisión (Año)"},
    {"value": "credito.cuotas", "label": "Crédito - Cantidad Cuotas"},
    {"value": "credito.tabla_transferencias", "label": "Crédito - Tabla de Transferencias"},
    {"value": "credito.tna_s_iva", "label": "Crédito - TNA (sin IVA)"},
    {"value": "credito.tea_s_iva", "label": "Crédito - TEA (sin IVA)"},
    {"value": "credito.tem_s_iva", "label": "Crédito - TEM (sin IVA)"},
    {"value": "credito.cft", "label": "Crédito - CFT Efectivo Anual"},
    {"value": "credito.fecha_emision", "label": "Crédito - Fecha Emisión"},
    {"value": "credito.primer_vto", "label": "Crédito - Primer Vencimiento"},
    {"value": "empresa.razon_social", "label": "Empresa - Razón Social"},
    {"value": "empresa.cuit", "label": "Empresa - CUIT"},
    {"value": "socio.razon_social", "label": "Socio Comercial - Razón Social"},
    {"value": "cartera.id", "label": "Cartera - ID"},
    {"value": "cartera.nombre", "label": "Cartera - Nombre"},
    {"value": "cartera.fecha_compra", "label": "Cartera - Fecha Venta/Compra"},
    {"value": "cartera.monto_total", "label": "Cartera - Monto Total"},
    {"value": "cartera.monto_total_letras", "label": "Cartera - Monto Total (En Letras)"},
    {"value": "cartera.socio_comercial", "label": "Cartera - Socio Comercial"},
    {"value": "cartera.recurso", "label": "Cartera - Con/Sin Recurso"},
    {"value": "cartera.fecha_dia", "label": "Cartera - Fecha Venta (Día)"},
    {"value": "cartera.fecha_mes_letras", "label": "Cartera - Fecha Venta (Mes Letras)"},
    {"value": "cartera.fecha_anio", "label": "Cartera - Fecha Venta (Año)"},
    {"value": "empresa.nro_cta_bancaria", "label": "Empresa - Nro. Cta. Bancaria"},
    {"value": "empresa.cbu", "label": "Empresa - CBU"},
    {"value": "empresa.domicilio", "label": "Empresa - Domicilio"},
    {"value": "empresa.nombre_banco", "label": "Empresa - Nombre del Banco"},
    {"value": "socio.cuit_formato", "label": "Socio Comercial - CUIT (con guiones)"},
    {"value": "socio.direccion", "label": "Socio Comercial - Dirección"},
    {"value": "socio.domicilio", "label": "Socio Comercial - Domicilio"},
    {"value": "socio.contacto", "label": "Socio Comercial - Contacto"},
    {"value": "cartera.valor_actual", "label": "Cartera - Valor Actual"},
    {"value": "cartera.valor_actual_letras", "label": "Cartera - Valor Actual (En Letras)"},
    {"value": "cartera.anexo_1", "label": "Cartera - Tabla Anexo A-I"},
    {"value": "cartera.anexo_2", "label": "Cartera - Tabla Anexo A-II"},
]

router = APIRouter(prefix="/api/v1/papeleria", tags=["Papeleria"])

DATA_DIR = "data/papeleria"

def _ensure_dir():
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR, exist_ok=True)

def _auto_map_variables(docx_path: str, doc_id: int, db: Session):
    try:
        from docx import Document
        import re
        doc = Document(docx_path)
    except Exception:
        return
        
    placeholders = set()
    pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')
    
    def _search_blocks(blocks):
        if hasattr(blocks, 'paragraphs'):
            for p in blocks.paragraphs:
                for match in pattern.findall(p.text):
                    placeholders.add(match.strip())
        if hasattr(blocks, 'tables'):
            for table in blocks.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if hasattr(cell, 'paragraphs'):
                            for p in cell.paragraphs:
                                for match in pattern.findall(p.text):
                                    placeholders.add(match.strip())
                                    
    _search_blocks(doc)
    
    for section in doc.sections:
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, header_type, None)
            if header:
                _search_blocks(header)
                
        for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_type, None)
            if footer:
                _search_blocks(footer)
                
    if not placeholders:
        return
        
    from src.database.models.papeleria import DocumentoVariable
    existing = db.query(DocumentoVariable).filter(DocumentoVariable.documento_id == doc_id).all()
    existing_placeholders = {v.placeholder for v in existing}
    
    new_vars = []
    for ph in placeholders:
        if ph not in existing_placeholders:
            new_vars.append(DocumentoVariable(documento_id=doc_id, placeholder=ph, system_field=""))
            
    if new_vars:
        db.add_all(new_vars)
        db.commit()

@router.post("/upload")
def upload_document(
    socio_id: int = Form(...),
    categoria: str = Form("creditos"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    # Validar que el socio existe si no es la empresa dueña (socio_id == 0)
    db_socio_id = socio_id if socio_id > 0 else None

    if db_socio_id is not None:
        socio = db.query(SocioComercial).filter(SocioComercial.id == db_socio_id).first()
        if not socio:
            raise HTTPException(status_code=404, detail="Socio comercial no encontrado.")

    # Validar extensión (Solo Word)
    filename = file.filename
    if not filename:
        raise HTTPException(status_code=400, detail="Nombre de archivo inválido.")
    
    ext = os.path.splitext(filename)[1].lower()
    if ext not in [".doc", ".docx"]:
        raise HTTPException(status_code=400, detail="Solo se permiten documentos Word (.doc, .docx).")

    _ensure_dir()
    
    # Para evitar colisiones, guardamos en data/papeleria/{db_socio_id o 0}_{filename}
    safe_filename = f"{db_socio_id or 0}_{filename}"
    file_path = os.path.join(DATA_DIR, safe_filename)

    # Si existe, lo sobreescribimos
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al guardar el archivo: {str(e)}")

    # Buscar si ya existe un registro para este socio y archivo y categoría
    if db_socio_id is not None:
        doc = db.query(DocumentoPapeleria).filter(
            DocumentoPapeleria.socio_id == db_socio_id,
            DocumentoPapeleria.nombre_archivo == filename,
            DocumentoPapeleria.categoria == categoria
        ).first()
    else:
        doc = db.query(DocumentoPapeleria).filter(
            DocumentoPapeleria.socio_id.is_(None),
            DocumentoPapeleria.nombre_archivo == filename,
            DocumentoPapeleria.categoria == categoria
        ).first()

    if doc:
        doc.ruta_archivo = file_path
        doc.tipo_archivo = ext
        from sqlalchemy.sql import func
        doc.fecha_subida = func.now()
    else:
        doc = DocumentoPapeleria(
            socio_id=db_socio_id,
            nombre_archivo=filename,
            ruta_archivo=file_path,
            tipo_archivo=ext,
            categoria=categoria
        )
        db.add(doc)

    db.commit()
    db.refresh(doc)

    if ext == ".docx":
        _auto_map_variables(file_path, doc.id, db)

    return {
        "status": "success",
        "message": "Documento subido correctamente",
        "documento": {
            "id": doc.id,
            "nombre_archivo": doc.nombre_archivo
        }
    }

@router.get("")
@router.get("/")
def list_documents(socio_id: int = None, categoria: str = None, db: Session = Depends(get_db)):
    query = db.query(DocumentoPapeleria)
    if socio_id is not None:
        if socio_id == 0:
            query = query.filter(DocumentoPapeleria.socio_id.is_(None))
        else:
            query = query.filter(DocumentoPapeleria.socio_id == socio_id)
            
    if categoria:
        query = query.filter(DocumentoPapeleria.categoria == categoria)
    
    docs = query.order_by(DocumentoPapeleria.orden.asc(), DocumentoPapeleria.fecha_subida.desc()).all()
    
    company = get_company_data(db)
    
    result = []
    for doc in docs:
        result.append({
            "id": doc.id,
            "socio_id": doc.socio_id if doc.socio_id is not None else 0,
            "socio_nombre": doc.socio.razon_social if doc.socio else company.razon_social,
            "nombre_archivo": doc.nombre_archivo,
            "tipo_archivo": doc.tipo_archivo,
            "categoria": doc.categoria,
            "orden": doc.orden,
            "fecha_subida": doc.fecha_subida.isoformat() if doc.fecha_subida else None
        })
    return result

from pydantic import BaseModel
class UpdateSocioPayload(BaseModel):
    socio_id: str | int | None

@router.patch("/{doc_id}/socio")
def update_document_socio(doc_id: int, payload: UpdateSocioPayload, db: Session = Depends(get_db)):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
        
    socio_id_val = None
    if str(payload.socio_id).strip() != "" and str(payload.socio_id) != "0":
        socio_id_val = int(payload.socio_id)
        
    if socio_id_val is not None:
        from src.database.models import Socio
        socio_exist = db.query(Socio).filter(Socio.id == socio_id_val).first()
        if not socio_exist:
            raise HTTPException(status_code=400, detail="El socio especificado no existe")

    doc.socio_id = socio_id_val
    db.commit()
    return {"status": "success", "message": "Socio actualizado correctamente"}

@router.get("/download/{doc_id}")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    
    if not os.path.exists(doc.ruta_archivo):
        raise HTTPException(status_code=404, detail="El archivo físico no existe en el servidor.")
    
    return FileResponse(
        path=doc.ruta_archivo,
        filename=doc.nombre_archivo,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@router.put("/{doc_id}/file")
def replace_document_file(
    doc_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    
    filename = file.filename
    if not filename:
        raise HTTPException(status_code=400, detail="Nombre de archivo inválido.")
    
    ext = os.path.splitext(filename)[1].lower()
    if ext not in [".doc", ".docx"]:
        raise HTTPException(status_code=400, detail="Solo se permiten documentos Word (.doc, .docx).")

    _ensure_dir()
    db_socio_id = doc.socio_id if doc.socio_id else 0
    safe_filename = f"{db_socio_id}_{filename}"
    file_path = os.path.join(DATA_DIR, safe_filename)

    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al guardar el archivo: {str(e)}")
        
    if doc.ruta_archivo != file_path and os.path.exists(doc.ruta_archivo):
        try:
            os.remove(doc.ruta_archivo)
        except:
            pass

    doc.ruta_archivo = file_path
    doc.nombre_archivo = filename
    doc.tipo_archivo = ext
    
    db.commit()
    
    if ext == ".docx":
        _auto_map_variables(file_path, doc.id, db)
        
    return {"status": "success", "message": "Archivo reemplazado correctamente"}

@router.delete("/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    
    if os.path.exists(doc.ruta_archivo):
        try:
            os.remove(doc.ruta_archivo)
        except Exception as e:
            # Continuamos aunque falle borrar el archivo físico (podría estar bloqueado)
            pass
            
    db.delete(doc)
    db.commit()
    
    return {"status": "success", "message": "Documento eliminado correctamente"}

@router.get("/system_fields")
def get_system_fields():
    return SYSTEM_FIELDS

@router.get("/{doc_id}/variables")
def get_document_variables(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    
    return [{"id": v.id, "placeholder": v.placeholder, "system_field": v.system_field} for v in doc.variables]

@router.post("/{doc_id}/variables")
def save_document_variables(doc_id: int, request: VariablesRequest, db: Session = Depends(get_db)):
    doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    
    # Delete existing
    db.query(DocumentoVariable).filter(DocumentoVariable.documento_id == doc_id).delete()
    
    # Add new
    for var_req in request.variables:
        new_var = DocumentoVariable(
            documento_id=doc_id,
            placeholder=var_req.placeholder,
            system_field=var_req.system_field
        )
        db.add(new_var)
    
    db.commit()
    return {"status": "success", "message": "Variables guardadas correctamente"}

@router.post("/reorder")
def reorder_documents(request: ReorderRequest, db: Session = Depends(get_db)):
    for item in request.documentos:
        doc = db.query(DocumentoPapeleria).filter(DocumentoPapeleria.id == item.id).first()
        if doc:
            doc.orden = item.orden
    
    db.commit()
    return {"status": "success", "message": "Orden actualizado correctamente"}

MESES = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

def resolve_system_field(credito: Credito, field: str):
    if field == "cliente.nombre":
        return f"{credito.cliente.nombre} {credito.cliente.apellido}"
    elif field == "cliente.id":
        return credito.cliente.cuil or ""
    elif field == "cliente.dni":
        return credito.cliente.documento
    elif field == "cliente.cuil":
        return credito.cliente.cuil
    elif field == "cliente.domicilio":
        return f"{credito.cliente.calle or ''} {credito.cliente.calle_nro or ''} {credito.cliente.localidad or ''}".strip()
    elif field == "cliente.cbu":
        return credito.cliente.cbu or ""
    elif field == "cliente.cuenta_bancaria":
        return credito.cliente.cuenta_bancaria or ""
    elif field == "cliente.banco":
        return credito.cliente.banco or ""
    elif field == "cliente.localidad":
        return credito.cliente.localidad or ""
    elif field == "cliente.provincia":
        return credito.cliente.provincia.nombre if credito.cliente.provincia else ""
    elif field == "cliente.nacionalidad":
        return credito.cliente.nacionalidad or ""
    elif field == "credito.monto_otorgado":
        return format_currency(credito.capital)
    elif field == "credito.monto_otorgado_letras":
        return monto_a_letras(float(credito.capital)) if credito.capital else monto_a_letras(0)
    elif field == "credito.monto_neto":
        if credito.capital:
            gastos = calcular_gastos_credito(credito)
            neto = float(credito.capital) - gastos
            return format_currency(neto)
        return "$ 0,00"
    elif field == "credito.monto_neto_letras":
        if credito.capital:
            gastos = calcular_gastos_credito(credito)
            neto = float(credito.capital) - gastos
            return monto_a_letras(neto)
        return monto_a_letras(0)
    elif field == "credito.gastos_otorgamiento":
        return format_currency(calcular_gastos_credito(credito))
    elif field == "credito.gastos_otorgamiento_letras":
        return monto_a_letras(calcular_gastos_credito(credito))
    elif field == "credito.sellado":
        if getattr(credito, 'comision', None) and getattr(credito.comision, 'porcentaje_sellado', 0):
            sellado = float(credito.capital) * float(credito.comision.porcentaje_sellado) / 100.0
            return format_currency(sellado)
        return "$ 0,00"
    elif field == "credito.sellado_letras":
        if getattr(credito, 'comision', None) and getattr(credito.comision, 'porcentaje_sellado', 0):
            sellado = float(credito.capital) * float(credito.comision.porcentaje_sellado) / 100.0
            return monto_a_letras(sellado)
        return monto_a_letras(0)
    elif field == "credito.interes":
        return format_currency(sum(float(c.interes or 0) for c in credito.cuotas)) if credito.cuotas else "$ 0,00"
    elif field == "credito.interes_letras":
        return monto_a_letras(sum(float(c.interes or 0) for c in credito.cuotas)) if credito.cuotas else monto_a_letras(0)
    elif field == "credito.iva":
        return format_currency(sum(float(c.iva or 0) for c in credito.cuotas)) if credito.cuotas else "$ 0,00"
    elif field == "credito.iva_letras":
        return monto_a_letras(sum(float(c.iva or 0) for c in credito.cuotas)) if credito.cuotas else monto_a_letras(0)
    elif field == "credito.fecha_alta":
        return credito.fecha_emision.strftime("%d/%m/%Y") if credito.fecha_emision else ""
    elif field == "credito.fecha_emision_dia":
        return str(credito.fecha_emision.day) if credito.fecha_emision else ""
    elif field == "credito.fecha_emision_mes_letras":
        if credito.fecha_emision:
            return MESES[credito.fecha_emision.month - 1]
        return ""
    elif field == "credito.fecha_emision_anio":
        return str(credito.fecha_emision.year) if credito.fecha_emision else ""
    elif field == "credito.cuotas":
        return str(len(credito.cuotas))
    elif field == "credito.id":
        return str(credito.id) if credito.id else ""
    elif field == "credito.id_externo":
        return str(credito.id_externo) if credito.id_externo else ""
    elif field == "cliente.apellido":
        return credito.cliente.apellido
    elif field == "cliente.estado_civil":
        return credito.cliente.estado_civil or ""
    elif field == "cliente.calle":
        return credito.cliente.calle or ""
    elif field == "cliente.calle_nro":
        return str(credito.cliente.calle_nro) if credito.cliente.calle_nro else ""
    elif field == "cliente.piso":
        return credito.cliente.piso or ""
    elif field == "cliente.depto":
        return credito.cliente.depto or ""
    elif field == "cliente.cp":
        return str(credito.cliente.id_codigo_postal) if credito.cliente.id_codigo_postal else ""
    elif field == "cliente.telefono":
        return credito.cliente.telefono or ""
    elif field == "cliente.telefono_2":
        return credito.cliente.telefono_2 or ""
    elif field == "cliente.mail":
        return credito.cliente.mail or ""
    elif field == "cliente.sexo":
        return credito.cliente.sexo.value if credito.cliente.sexo else ""
    elif field == "cliente.fecha_nacimiento":
        return credito.cliente.fecha_nacimiento.strftime("%d/%m/%Y") if credito.cliente.fecha_nacimiento else ""
    elif field == "cliente.cargo":
        return credito.cliente.cargo or ""
    elif field == "cliente.pep":
        return "SÍ" if credito.cliente.pep else "NO"
    elif field == "cliente.repet":
        return "SÍ" if credito.cliente.repet else "NO"
    elif field in ["cliente.repet_id", "cliente.repet_fecha_consulta", "cliente.repet_estado"]:
        from src.database.models.repet import RepetAuditLog
        from sqlalchemy.orm import object_session
        db = object_session(credito)
        if not db:
            return ""
        full_name = f"{credito.cliente.nombre} {credito.cliente.apellido}".strip()
        last_log = db.query(RepetAuditLog).filter(
            RepetAuditLog.searched_name == full_name
        ).order_by(RepetAuditLog.timestamp.desc()).first()
        
        if not last_log:
            if field == "cliente.repet_estado":
                return "SIN CONSULTA"
            return ""
            
        if field == "cliente.repet_id":
            return str(last_log.id)
        elif field == "cliente.repet_fecha_consulta":
            if not last_log.timestamp:
                return ""
            try:
                import pytz
                local_tz = pytz.timezone('America/Argentina/Buenos_Aires')
                ts = last_log.timestamp
                if ts.tzinfo is None:
                    ts = pytz.utc.localize(ts)
                return ts.astimezone(local_tz).strftime("%d/%m/%Y %H:%M:%S")
            except Exception:
                return last_log.timestamp.strftime("%d/%m/%Y %H:%M:%S")
        elif field == "cliente.repet_estado":
            return "POSITIVO" if getattr(last_log, 'is_match', False) else "NEGATIVO"
    elif field == "empleador.razon_social":
        return credito.cliente.empleador.razon_social if credito.cliente.empleador else ""
    elif field == "empleador.fecha_ingreso":
        return credito.cliente.fecha_ingreso.strftime("%d/%m/%Y") if credito.cliente.fecha_ingreso else ""
    elif field == "empleador.ingreso_mensual":
        return format_currency(credito.cliente.remuneracion) if credito.cliente.remuneracion else "$ 0,00"
    elif field == "empleador.legajo":
        return credito.cliente.legajo or ""
    elif field == "empleador.domicilio_calle":
        return credito.cliente.empleador.domicilio_calle if credito.cliente.empleador else ""
    elif field == "empleador.domicilio_nro":
        return str(credito.cliente.empleador.domicilio_nro) if credito.cliente.empleador and credito.cliente.empleador.domicilio_nro else ""
    elif field == "empleador.domicilio_piso":
        return credito.cliente.empleador.domicilio_piso if credito.cliente.empleador else ""
    elif field == "empleador.domicilio_depto":
        return credito.cliente.empleador.domicilio_depto if credito.cliente.empleador else ""
    elif field == "empleador.localidad":
        return credito.cliente.empleador.localidad if credito.cliente.empleador else ""
    elif field == "empleador.provincia":
        return credito.cliente.empleador.provincia.nombre if credito.cliente.empleador and credito.cliente.empleador.provincia else ""
    elif field == "empleador.codigo_postal":
        return str(credito.cliente.empleador.id_codigo_postal) if credito.cliente.empleador and credito.cliente.empleador.id_codigo_postal else ""
    elif field == "empleador.telefono":
        return credito.cliente.empleador.telefono if credito.cliente.empleador else ""
    elif field.startswith("referido_"):
        parts = field.split(".")
        if len(parts) == 2:
            try:
                idx = int(parts[0].split("_")[1]) - 1
                if 0 <= idx < len(credito.cliente.referidos):
                    ref = credito.cliente.referidos[idx]
                    return getattr(ref, parts[1], "") or ""
            except:
                pass
        return ""
    elif field == "credito.plazo":
        return str(credito.plazo)
    elif field == "credito.valor_cuota":
        if credito.cuotas:
            return format_currency(float(credito.cuotas[0].capital or 0) + float(credito.cuotas[0].interes or 0) + float(credito.cuotas[0].iva or 0))
        return "$ 0,00"
    elif field == "credito.monto_total":
        if credito.cuotas:
            return format_currency(sum(float(c.capital or 0) + float(c.interes or 0) + float(c.iva or 0) for c in credito.cuotas))
        return "$ 0,00"
    elif field == "credito.monto_total_letras":
        if credito.cuotas:
            return monto_a_letras(sum(float(c.capital or 0) + float(c.interes or 0) + float(c.iva or 0) for c in credito.cuotas))
        return monto_a_letras(0)
    elif field == "credito.tna_c_iva":
        return str(credito.tna_c_iva)
    elif field == "credito.tna_s_iva":
        if credito.tna_c_iva is not None:
            tna_s_iva = float(credito.tna_c_iva) / 1.21
            return f"{(tna_s_iva * 100):.2f} %"
        return "N/A"
    elif field == "credito.tea_s_iva":
        if credito.tna_c_iva is not None:
            tna_s_iva = float(credito.tna_c_iva) / 1.21
            tea_s_iva = (1 + tna_s_iva * 30 / 365) ** (365 / 30) - 1
            return f"{(tea_s_iva * 100):.2f} %"
        return "N/A"
    elif field == "credito.tem_s_iva":
        if credito.tna_c_iva is not None:
            tna_s_iva = float(credito.tna_c_iva) / 1.21
            tea_s_iva = (1 + tna_s_iva * 30 / 365) ** (365 / 30) - 1
            tem_s_iva = (1 + tea_s_iva) ** (30 / 365) - 1
            return f"{(tem_s_iva * 100):.2f} %"
        return "N/A"
    elif field == "credito.cft":
        if credito.capital and credito.cuotas:
            gastos = calcular_gastos_credito(credito)
            cf_0 = - (float(credito.capital) - gastos)
            flujos = [cf_0]
            for c in credito.cuotas:
                flujos.append(float(c.capital or 0) + float(c.interes or 0) + float(c.iva or 0))
            tir = npf.irr(flujos)
            if not isinstance(tir, float) or tir is None:
                return "N/A"
            cft_tea = (1 + tir) ** (365 / 30) - 1
            return f"{(cft_tea * 100):.2f} %"
        return "N/A"
    elif field == "credito.fecha_emision":
        return credito.fecha_emision.strftime("%d/%m/%Y") if credito.fecha_emision else ""
    elif field == "credito.primer_vto":
        if credito.cuotas and len(credito.cuotas) > 0:
            return credito.cuotas[0].fecha_vencimiento.strftime("%d/%m/%Y") if credito.cuotas[0].fecha_vencimiento else ""
        return ""
    elif field == "empresa.razon_social":
        return get_company_data(db).razon_social
    elif field == "empresa.cuit":
        return get_company_data(db).cuit
    elif field == "socio.razon_social":
        return credito.socio_originador.razon_social if credito.socio_originador else ""
    elif field == "credito.tabla_transferencias":
        transferencias_data = []
        for t in credito.transferencias:
            transferencias_data.append({
                "CUIL/CUIT": t.cuit,
                "CBU": t.cbu,
                "Razón Social": t.razon_social,
                "Monto": format_currency(t.monto)
            })
        return {"__type__": "table", "headers": ["CUIL/CUIT", "CBU", "Razón Social", "Monto"], "rows": transferencias_data}
    return ""

def resolve_system_field_cartera(cartera: Cartera, field: str, db: Session = None):
    if field == "cartera.id":
        return str(cartera.id)
    elif field == "cartera.nombre":
        return cartera.nombre or ""
    elif field == "cartera.fecha_compra":
        return cartera.fecha_compra.strftime("%d/%m/%Y") if cartera.fecha_compra else ""
    elif field == "cartera.monto_total":
        total = sum(float(c.capital or 0) for c in cartera.creditos_incluidos) if cartera.creditos_incluidos else 0
        return format_currency(total)
    elif field == "cartera.monto_total_letras":
        total = sum(float(c.capital or 0) for c in cartera.creditos_incluidos) if cartera.creditos_incluidos else 0
        return monto_a_letras(total)
    elif field == "cartera.socio_comercial" or field == "socio.razon_social":
        return cartera.socio.razon_social if cartera.socio else ""
    elif field == "empresa.razon_social":
        return get_company_data(db).razon_social
    elif field == "empresa.cuit":
        return get_company_data(db).cuit
    elif field == "cartera.recurso":
        return "CON RECURSO" if cartera.recurso else "SIN RECURSO"
    elif field == "cartera.fecha_dia":
        return str(cartera.fecha_compra.day) if cartera.fecha_compra else ""
    elif field == "cartera.fecha_mes_letras":
        if cartera.fecha_compra:
            meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
            return meses[cartera.fecha_compra.month - 1]
        return ""
    elif field == "cartera.fecha_anio":
        return str(cartera.fecha_compra.year) if cartera.fecha_compra else ""
    elif field == "empresa.nro_cta_bancaria":
        return get_company_data(db).bank_account
    elif field == "empresa.cbu":
        return get_company_data(db).cbu
    elif field == "empresa.domicilio":
        return get_company_data(db).domicilio
    elif field == "empresa.nombre_banco":
        return get_company_data(db).bank_name
    elif field == "socio.cuit_formato":
        cuit = cartera.socio.cuit if cartera.socio else ""
        if len(cuit) == 11:
            return f"{cuit[:2]}-{cuit[2:10]}-{cuit[10:]}"
        return cuit
    elif field == "socio.direccion" or field == "socio.domicilio":
        return cartera.socio.domicilio_legal if cartera.socio else ""
    elif field == "socio.contacto":
        return cartera.socio.contacto_nombre if cartera.socio else ""
    elif field in ("cartera.valor_actual", "cartera.valor_actual_letras"):
        tna = float(cartera.tna_descuento) if cartera.tna_descuento else 0.0
        va_total = 0.0
        if cartera.operaciones:
            for op in cartera.operaciones:
                if op.cuota_comercializada and op.cuota:
                    c = op.cuota
                    total_cuota = float(c.capital or 0) + float(c.interes or 0)
                    if c.fecha_vencimiento and cartera.fecha_compra:
                        dias = max(0, (c.fecha_vencimiento - cartera.fecha_compra).days)
                        va_cuota = total_cuota / ((1 + (tna * 30 / 365)) ** (dias / 30))
                    else:
                        va_cuota = total_cuota
                    va_total += va_cuota
        
        if field == "cartera.valor_actual":
            return format_currency(va_total)
        else:
            return monto_a_letras(va_total)
    
    elif field == "cartera.anexo_1":
        tna = float(cartera.tna_descuento) if cartera.tna_descuento else 0.0
        headers = ['OPERACIÓN', 'APELLIDO', 'NOMBRE', 'Tipo Doc.', 'Nro Doc.', 'LINEA', 'ORGANISMO', 'FECHA LIQUIDACION', 'PLAZO OPERACIÓN', 'PLAZO VENDIDO', 'VALOR NOMINAL', 'CAPITAL + INTERES', 'VALOR ACTUAL', 'TASA']
        rows = []
        tot_vn = 0.0
        tot_cap_int = 0.0
        tot_va = 0.0
        credito_ops = {}
        for op in cartera.operaciones:
            if not op.cuota_comercializada or not op.cuota:
                continue
            cid = op.cuota.credito_id
            if cid not in credito_ops:
                credito_ops[cid] = []
            credito_ops[cid].append(op)
        for cid, ops in credito_ops.items():
            credito = ops[0].cuota.credito
            cli = credito.cliente
            plazo_vend = len(ops)
            vn = sum(float(op.cuota.capital or 0) for op in ops)
            cap_int = sum(float(op.cuota.capital or 0) + float(op.cuota.interes or 0) + float(op.cuota.iva or 0) for op in ops)
            va_credito = 0.0
            for op in ops:
                c = op.cuota
                total_c = float(c.capital or 0) + float(c.interes or 0)
                if c.fecha_vencimiento and cartera.fecha_compra:
                    dias = max(0, (c.fecha_vencimiento - cartera.fecha_compra).days)
                    va_credito += total_c / ((1 + (tna * 30 / 365)) ** (dias / 30))
                else:
                    va_credito += total_c
            tot_vn += vn
            tot_cap_int += cap_int
            tot_va += va_credito
            rows.append({
                "OPERACIÓN": str(credito.id_externo or credito.id),
                "APELLIDO": cli.apellido if cli else "",
                "NOMBRE": cli.nombre if cli else "",
                "Tipo Doc.": "DNI" if cli else "",
                "Nro Doc.": cli.documento if cli else "",
                "LINEA": credito.socio_originador.razon_social if credito.socio_originador else "",
                "ORGANISMO": cli.empleador.razon_social if cli and cli.empleador else "",
                "FECHA LIQUIDACION": credito.fecha_emision.strftime("%d/%m/%Y") if credito.fecha_emision else "",
                "PLAZO OPERACIÓN": str(credito.plazo or ""),
                "PLAZO VENDIDO": str(plazo_vend),
                "VALOR NOMINAL": format_currency(vn),
                "CAPITAL + INTERES": format_currency(cap_int),
                "VALOR ACTUAL": format_currency(va_credito),
                "TASA": f"{tna:.2f}%"
            })
        rows.append({
            "OPERACIÓN": "TOTALES", "APELLIDO": "", "NOMBRE": "", "Tipo Doc.": "", "Nro Doc.": "",
            "LINEA": "", "ORGANISMO": "", "FECHA LIQUIDACION": "", "PLAZO OPERACIÓN": "", "PLAZO VENDIDO": "",
            "VALOR NOMINAL": format_currency(tot_vn), "CAPITAL + INTERES": format_currency(tot_cap_int),
            "VALOR ACTUAL": format_currency(tot_va), "TASA": ""
        })
        return {"__type__": "table", "headers": headers, "rows": rows}
    
    elif field == "cartera.anexo_2":
        tna = float(cartera.tna_descuento) if cartera.tna_descuento else 0.0
        headers = ['Fecha de Vencimiento', 'Fecha de Rendición', 'Capital', 'Interés', 'Monto Financiado', 'Tasa', 'Valor Actual', 'Fecha de Cesión', 'Entidad Cesionaria', 'Nº Cesión']
        rows = []
        from collections import defaultdict
        due_dates = defaultdict(list)
        for op in cartera.operaciones:
            if op.cuota_comercializada and op.cuota and op.cuota.fecha_vencimiento:
                due_dates[op.cuota.fecha_vencimiento].append(op.cuota)
        tot_cap = 0.0
        tot_int = 0.0
        tot_monto = 0.0
        tot_va = 0.0
        for fv in sorted(due_dates.keys()):
            cuotas_fv = due_dates[fv]
            import calendar
            last_day = calendar.monthrange(fv.year, fv.month)[1]
            fr = fv.replace(day=last_day)
            cap = sum(float(c.capital or 0) for c in cuotas_fv)
            int = sum(float(c.interes or 0) for c in cuotas_fv)
            monto = cap + int
            va_fv = 0.0
            if cartera.fecha_compra:
                dias = max(0, (fv - cartera.fecha_compra).days)
                va_fv = monto / ((1 + (tna * 30 / 365)) ** (dias / 30))
            else:
                va_fv = monto
            tot_cap += cap
            tot_int += int
            tot_monto += monto
            tot_va += va_fv
            rows.append({
                "Fecha de Vencimiento": fv.strftime("%d/%m/%Y"),
                "Fecha de Rendición": fr.strftime("%d/%m/%Y"),
                "Capital": format_currency(cap),
                "Interés": format_currency(int),
                "Monto Financiado": format_currency(monto),
                "Tasa": f"{tna:.2%}",
                "Valor Actual": format_currency(va_fv),
                "Fecha de Cesión": cartera.fecha_compra.strftime("%d/%m/%Y") if cartera.fecha_compra else "",
                "Entidad Cesionaria": cartera.socio.razon_social if cartera.socio else "",
                "Nº Cesión": str(cartera.id).zfill(3)
            })
        rows.append({
            "Fecha de Vencimiento": "TOTALES", "Fecha de Rendición": "",
            "Capital": format_currency(tot_cap), "Interés": format_currency(tot_int),
            "Monto Financiado": format_currency(tot_monto), "Tasa": "",
            "Valor Actual": format_currency(tot_va), "Fecha de Cesión": "",
            "Entidad Cesionaria": "", "Nº Cesión": ""
        })
        return {"__type__": "table", "headers": headers, "rows": rows}
    
    return ""

@router.post("/generar_por_credito/{credito_id}")
def generar_papeleria_credito(credito_id: int, db: Session = Depends(get_db)):
    credito = db.query(Credito).filter(Credito.id == credito_id).first()
    if not credito:
        raise HTTPException(status_code=404, detail="Crédito no encontrado.")
    
    final_pdf_path = _generar_pdf_for_credito(credito, db)

    return FileResponse(
        path=final_pdf_path,
        filename=f"Legajo_Credito_{credito_id}.pdf",
        media_type='application/pdf'
    )

def _generar_pdf_for_credito(credito: Credito, db: Session) -> str:
    socio_ids = [None]
    
    # Extraer socios de la tasa y comisión
    if credito.comision:
        if credito.comision.socio_originador_id:
            socio_ids.append(credito.comision.socio_originador_id)
        if credito.comision.socio_intermediario_id:
            socio_ids.append(credito.comision.socio_intermediario_id)
        if credito.comision.gasto_1_socio_id:
            socio_ids.append(credito.comision.gasto_1_socio_id)
        if credito.comision.gasto_2_socio_id:
            socio_ids.append(credito.comision.gasto_2_socio_id)

    # Si por alguna razón tiene originador directo pero no está en la lista (caso legado)
    if credito.socio_originador_id and credito.socio_originador_id not in socio_ids:
        socio_ids.append(credito.socio_originador_id)

    # Filtro combinado porque in_() con SQL no funciona para NULL
    from sqlalchemy import or_
    valid_ids = [s for s in socio_ids if s is not None]
    
    docs = db.query(DocumentoPapeleria).filter(
        DocumentoPapeleria.categoria == "creditos",
        or_(
            DocumentoPapeleria.socio_id.is_(None),
            DocumentoPapeleria.socio_id.in_(valid_ids) if valid_ids else False
        )
    ).order_by(
        DocumentoPapeleria.socio_id.isnot(None),
        DocumentoPapeleria.orden.asc()
    ).all()

    if not docs:
        raise HTTPException(status_code=400, detail="No hay documentos de papelería configurados para esta operación.")

    import tempfile
    from pypdf import PdfWriter
    from src.logic.creditos.legajos import process_document

    merger = PdfWriter()
    temp_files = []
    output_dir = "data/legajos"
    os.makedirs(output_dir, exist_ok=True)
    
    # Si el crédito no tiene id aún (simulación) le ponemos un prefijo random
    credito_ident = credito.id if credito.id else f"simulacion_{credito.cliente_cuil}"
    final_pdf_path = os.path.join(output_dir, f"credito_{credito_ident}.pdf")

    has_pythoncom = False
    try:
        import pythoncom
        has_pythoncom = True
    except ImportError:
        pass

    try:
        try:
            if has_pythoncom:
                pythoncom.CoInitialize()
            for doc in docs:
                data = {}
                for var in doc.variables:
                    data[var.placeholder] = resolve_system_field(credito, var.system_field)

                fd, temp_pdf = tempfile.mkstemp(suffix=".pdf")
                os.close(fd)
                os.remove(temp_pdf)
                temp_files.append(temp_pdf)
                
                process_document(doc.ruta_archivo, temp_pdf, data)
                merger.append(temp_pdf)

            merger.write(final_pdf_path)
            merger.close()

        finally:
            if has_pythoncom:
                pythoncom.CoUninitialize()
            for f in temp_files:
                if os.path.exists(f):
                    try:
                        os.remove(f)
                    except:
                        pass
    except Exception as e:
        import traceback
        error_msg = f"Error interno: {str(e)}\n\n{traceback.format_exc()}"
        raise HTTPException(status_code=500, detail=error_msg)
        
    return final_pdf_path

@router.post("/generar_por_cartera/{cartera_id}")
def generar_papeleria_cartera(cartera_id: int, formato: str = 'pdf', db: Session = Depends(get_db)):
    cartera = db.query(Cartera).filter(Cartera.id == cartera_id).first()
    if not cartera:
        raise HTTPException(status_code=404, detail="Cartera no encontrada.")
    
    socio_id = cartera.socio_id
    
    from sqlalchemy import or_
    docs = db.query(DocumentoPapeleria).filter(
        DocumentoPapeleria.categoria == "ventas_cartera",
        or_(
            DocumentoPapeleria.socio_id.is_(None),
            DocumentoPapeleria.socio_id == socio_id if socio_id else False
        )
    ).order_by(
        DocumentoPapeleria.socio_id.isnot(None),
        DocumentoPapeleria.orden.asc()
    ).all()
    
    if not docs:
        raise HTTPException(status_code=400, detail="No hay documentos de papelería de ventas de cartera configurados para este socio (o genéricos).")
        
    import tempfile
    from src.logic.creditos.legajos import process_document, process_docx
    from pypdf import PdfWriter
    
    has_pythoncom = False
    try:
        import pythoncom
        has_pythoncom = True
    except ImportError:
        pass
    
    output_dir = "data/legajos"
    os.makedirs(output_dir, exist_ok=True)
    
    if formato == 'docx' and len(docs) == 1:
        doc = docs[0]
        data = {}
        for var in doc.variables:
            data[var.placeholder] = resolve_system_field_cartera(cartera, var.system_field, db)
            
        final_docx_path = os.path.join(output_dir, f"Contrato_Cartera_{cartera_id}.docx")
        process_docx(doc.ruta_archivo, final_docx_path, data)
        return FileResponse(
            path=final_docx_path,
            filename=f"Contrato_Cartera_{cartera_id}.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    # If multiple docs or forced PDF, merge as PDF
    if has_pythoncom:
        try:
            pythoncom.CoInitialize()
        except Exception:
            has_pythoncom = False
        
    merger = PdfWriter()
    temp_files = []
    final_pdf_path = os.path.join(output_dir, f"Contrato_Cartera_{cartera_id}.pdf")
    
    try:
        try:
            for doc in docs:
                data = {}
                for var in doc.variables:
                    data[var.placeholder] = resolve_system_field_cartera(cartera, var.system_field, db)

                fd, temp_pdf = tempfile.mkstemp(suffix=".pdf")
                os.close(fd)
                os.remove(temp_pdf)
                temp_files.append(temp_pdf)
                
                process_document(doc.ruta_archivo, temp_pdf, data)
                merger.append(temp_pdf)

            merger.write(final_pdf_path)
            return FileResponse(
                path=final_pdf_path,
                filename=f"Contrato_Cartera_{cartera_id}.pdf",
                media_type='application/pdf'
            )
        finally:
            if has_pythoncom:
                pythoncom.CoUninitialize()
            for t in temp_files:
                if os.path.exists(t):
                    try:
                        os.remove(t)
                    except:
                        pass
    except Exception as e:
        import traceback
        error_msg = f"Error interno: {str(e)}\n\n{traceback.format_exc()}"
        raise HTTPException(status_code=500, detail=error_msg)
