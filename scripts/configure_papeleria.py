import os
import re
import sys
from docx import Document

# Add the project root to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from src.database import get_db
from src.database.models.papeleria import DocumentoPapeleria, DocumentoVariable
from src.api.routes.papeleria import SYSTEM_FIELDS

def extract_placeholders_from_docx(filepath):
    placeholders = set()
    pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')
    
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return placeholders
        
    doc = Document(filepath)
    
    def process_blocks(blocks):
        if hasattr(blocks, 'paragraphs'):
            for p in blocks.paragraphs:
                for match in pattern.findall(p.text):
                    placeholders.add(match.strip())
        if hasattr(blocks, 'tables'):
            for table in blocks.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if hasattr(cell, 'paragraphs'):
                            for p in cell.paragraphs:
                                for match in pattern.findall(p.text):
                                    placeholders.add(match.strip())
                                    
    process_blocks(doc)
    for section in doc.sections:
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, header_type, None)
            if header:
                process_blocks(header)
        for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_type, None)
            if footer:
                process_blocks(footer)
                
    return placeholders

def configure_papeleria():
    db = next(get_db())
    
    valid_fields = {f['value'] for f in SYSTEM_FIELDS}
    
    documents = db.query(DocumentoPapeleria).all()
    print(f"Found {len(documents)} documents.")
    
    for doc in documents:
        print(f"\nProcessing {doc.nombre_archivo}...")
        
        # Clear existing variables
        db.query(DocumentoVariable).filter(DocumentoVariable.documento_id == doc.id).delete()
        
        placeholders = extract_placeholders_from_docx(doc.ruta_archivo)
        print(f"Found placeholders: {placeholders}")
        
        custom_mapping = {
            'credito_fecha_emision_dia': 'credito.fecha_emision_dia',
            'credito_fecha_emision_mes_letras': 'credito.fecha_emision_mes_letras',
            'credito_fecha_emision_anio': 'credito.fecha_emision_anio',
            'cliente_nombre_completo': 'cliente.nombre',
            'credito_nro': 'credito.id',
            'cliente_telefono_1': 'cliente.telefono',
            'cliente_telefono_2': 'cliente.telefono_2',
            'cliente_calle': 'cliente.calle',
            'cliente_calle_nro': 'cliente.calle_nro',
            'importe_capital': 'credito.monto_otorgado',
            'monto_total': 'credito.monto_total',
            'plazo_meses': 'credito.plazo',
            'tasa_interes_compensatorio': 'credito.tna_c_iva',
            'empleador_ingreso_mensual': 'empleador.ingreso_mensual',
            'empleador_fecha_ingreso': 'empleador.fecha_ingreso',
            'empleador_nombre': 'empleador.razon_social',
            'cliente_email': 'cliente.mail',
            'fecha_nacimiento': 'cliente.fecha_nacimiento',
        }
        
        for p in placeholders:
            # Try to find a matching system field
            system_field = None
            
            # Check custom mapping
            if p in custom_mapping and custom_mapping[p] in valid_fields:
                system_field = custom_mapping[p]
            # Check exact match
            if p in valid_fields:
                system_field = p
            else:
                # Check replacing _ with .
                dot_version = p.replace('_', '.')
                if dot_version in valid_fields:
                    system_field = dot_version
                else:
                    # Basic guessing
                    for f in valid_fields:
                        if p.lower() in f.lower() or dot_version.lower() in f.lower():
                            system_field = f
                            break
            
            if not system_field:
                print(f"  Warning: No matching system field for '{p}', defaulting to itself.")
                system_field = p
            else:
                print(f"  Mapped '{p}' to '{system_field}'")
                
            var = DocumentoVariable(
                documento_id=doc.id,
                placeholder=p,
                system_field=system_field
            )
            db.add(var)
            
    db.commit()
    print("\nAll done!")

if __name__ == "__main__":
    configure_papeleria()
